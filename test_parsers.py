#!/usr/bin/env python3
"""
Test script for document parsers and chart generator
"""
import os
import sys
import django
from pathlib import Path

# Add the project root to Python path
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

# Setup Django
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'chatbot.settings.development')
django.setup()

def test_imports():
    """Test that all imports work correctly"""
    print("Testing imports...")
    
    try:
        # Test document parsers
        from apps.documents.parsers import PDFParser, PDFContent
        from apps.documents.parsers import ExcelParser, ExcelContent
        from apps.documents.parsers import WordParser, WordContent
        print("✓ Document parsers imported successfully")
        
        # Test chart generator
        from apps.agents.tools.chart_generator import ChartGenerator
        print("✓ Chart generator imported successfully")
        
        # Test that required dependencies are available
        import pdfplumber
        import PyPDF2
        import pandas as pd
        import matplotlib.pyplot as plt
        from docx import Document
        print("✓ All required dependencies available")
        
        return True
        
    except ImportError as e:
        print(f"✗ Import error: {e}")
        return False
    except Exception as e:
        print(f"✗ Unexpected error: {e}")
        return False

def create_sample_excel():
    """Create a sample Excel file for testing"""
    import pandas as pd
    
    # Create sample data
    data = {
        'Name': ['Alice', 'Bob', 'Charlie', 'Diana'],
        'Age': [25, 30, 35, 28],
        'Salary': [50000, 60000, 70000, 55000],
        'Department': ['IT', 'HR', 'Finance', 'IT']
    }
    
    df = pd.DataFrame(data)
    
    # Create test directory
    test_dir = Path('test_files')
    test_dir.mkdir(exist_ok=True)
    
    file_path = test_dir / 'sample.xlsx'
    df.to_excel(file_path, index=False)
    
    return str(file_path)

def create_sample_word():
    """Create a sample Word document for testing"""
    from docx import Document
    
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph for the Word parser.')
    doc.add_heading('Section 1', level=1)
    doc.add_paragraph('This is another paragraph with some content.')
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Header 1'
    table.cell(0, 1).text = 'Header 2'
    table.cell(1, 0).text = 'Data 1'
    table.cell(1, 1).text = 'Data 2'
    
    test_dir = Path('test_files')
    test_dir.mkdir(exist_ok=True)
    
    file_path = test_dir / 'sample.docx'
    doc.save(file_path)
    
    return str(file_path)

def test_excel_parser():
    """Test Excel parser functionality"""
    print("\nTesting Excel parser...")
    
    try:
        from apps.documents.parsers import ExcelParser
        
        # Create sample file
        file_path = create_sample_excel()
        print(f"Created sample Excel file: {file_path}")
        
        # Test parsing
        content = ExcelParser.parse(file_path)
        print(f"✓ Parsed Excel file with {content.metadata['sheet_count']} sheets")
        print(f"  Total rows: {content.metadata['total_rows']}")
        print(f"  Total columns: {content.metadata['total_columns']}")
        
        # Test summary generation
        summary = ExcelParser.generate_summary(content)
        print(f"✓ Generated summary: {summary[:100]}...")
        
        return True
        
    except Exception as e:
        print(f"✗ Excel parser error: {e}")
        return False

def test_word_parser():
    """Test Word parser functionality"""
    print("\nTesting Word parser...")
    
    try:
        from apps.documents.parsers import WordParser
        
        # Create sample file
        file_path = create_sample_word()
        print(f"Created sample Word file: {file_path}")
        
        # Test parsing
        content = WordParser.parse(file_path)
        print(f"✓ Parsed Word file with {content.metadata['paragraph_count']} paragraphs")
        print(f"  Word count: {content.metadata['word_count']}")
        print(f"  Headers count: {content.metadata['header_count']}")
        
        # Test summary generation
        summary = WordParser.generate_summary(content)
        print(f"✓ Generated summary: {summary[:100]}...")
        
        return True
        
    except Exception as e:
        print(f"✗ Word parser error: {e}")
        return False

def test_chart_generator():
    """Test chart generator functionality"""
    print("\nTesting Chart generator...")
    
    try:
        from apps.agents.tools.chart_generator import ChartGenerator
        
        # Test data
        data = {
            'x': ['A', 'B', 'C', 'D'],
            'y': [10, 20, 15, 25]
        }
        
        # Test bar chart
        chart_path = ChartGenerator.generate_chart(
            data=data,
            chart_type='bar',
            title='Test Bar Chart',
            xlabel='Categories',
            ylabel='Values'
        )
        print(f"✓ Generated bar chart: {chart_path}")
        
        # Test line chart
        chart_path2 = ChartGenerator.generate_chart(
            data=data,
            chart_type='line',
            title='Test Line Chart'
        )
        print(f"✓ Generated line chart: {chart_path2}")
        
        return True
        
    except Exception as e:
        print(f"✗ Chart generator error: {e}")
        return False

def test_error_handling():
    """Test error handling with invalid files"""
    print("\nTesting error handling...")
    
    try:
        from apps.documents.parsers import PDFParser, ExcelParser, WordParser
        from apps.agents.tools.chart_generator import ChartGenerator
        
        # Test with non-existent file
        try:
            ExcelParser.parse("non_existent_file.xlsx")
            print("✗ Should have raised an error for non-existent file")
            return False
        except Exception:
            print("✓ Correctly handled non-existent file error")
        
        # Test chart generator with invalid chart type
        try:
            ChartGenerator.generate_chart(
                data={'x': [1, 2], 'y': [3, 4]},
                chart_type='invalid_type'
            )
            print("✗ Should have raised an error for invalid chart type")
            return False
        except ValueError:
            print("✓ Correctly handled invalid chart type error")
        
        return True
        
    except Exception as e:
        print(f"✗ Error handling test failed: {e}")
        return False

def main():
    """Run all tests"""
    print("🧪 Testing Ultra PDF Chatbot 3000 Document Parsing Tools")
    print("=" * 60)
    
    results = []
    
    # Run tests
    results.append(("Imports", test_imports()))
    results.append(("Excel Parser", test_excel_parser()))
    results.append(("Word Parser", test_word_parser()))
    results.append(("Chart Generator", test_chart_generator()))
    results.append(("Error Handling", test_error_handling()))
    
    # Summary
    print("\n" + "=" * 60)
    print("TEST RESULTS:")
    print("=" * 60)
    
    passed = 0
    for test_name, result in results:
        status = "✓ PASS" if result else "✗ FAIL"
        print(f"{test_name}: {status}")
        if result:
            passed += 1
    
    print(f"\nTotal: {passed}/{len(results)} tests passed")
    
    if passed == len(results):
        print("🎉 All tests passed! Document processing tools are ready.")
    else:
        print("❌ Some tests failed. Please check the errors above.")
    
    return passed == len(results)

if __name__ == '__main__':
    success = main()
    sys.exit(0 if success else 1)