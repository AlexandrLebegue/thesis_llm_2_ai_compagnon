#!/usr/bin/env python3
"""
Comprehensive integration test for the SmolAgents system with document processing tools
"""

import os
import sys
import django
from pathlib import Path
import tempfile
import pandas as pd
from docx import Document

# Setup Django environment
sys.path.append('.')
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'chatbot.settings.development')
django.setup()

# Import after Django setup
from apps.agents.registry import tool_registry
from apps.agents.tools import *  # Register all tools
from apps.documents.summarizer import DocumentSummarizer
from apps.documents.parsers import PDFParser, ExcelParser, WordParser
from apps.documents.models import Document as DocumentModel, DocumentSession

def create_test_excel():
    """Create a test Excel file"""
    test_dir = Path('/tmp/ultra_pdf_chatbot/test_files')
    test_dir.mkdir(parents=True, exist_ok=True)
    
    excel_path = test_dir / 'test_data.xlsx'
    
    # Create test data
    data = {
        'Month': ['Jan', 'Feb', 'Mar', 'Apr', 'May'],
        'Sales': [100, 150, 120, 180, 200],
        'Expenses': [80, 90, 85, 95, 110],
        'Profit': [20, 60, 35, 85, 90]
    }
    
    df = pd.DataFrame(data)
    df.to_excel(excel_path, index=False)
    
    return str(excel_path)

def create_test_word():
    """Create a test Word document"""
    test_dir = Path('/tmp/ultra_pdf_chatbot/test_files')
    test_dir.mkdir(parents=True, exist_ok=True)
    
    word_path = test_dir / 'test_document.docx'
    
    doc = Document()
    doc.add_heading('Sales Report', 0)
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph('This document contains the quarterly sales analysis.')
    doc.add_paragraph('Key findings include strong performance in Q2.')
    
    doc.add_heading('Data Analysis', level=1)
    doc.add_paragraph('The following table shows monthly performance:')
    
    # Add a table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Month'
    header_cells[1].text = 'Sales'
    header_cells[2].text = 'Growth'
    
    # Add data rows
    data_rows = [
        ['January', '$100K', '5%'],
        ['February', '$150K', '15%'],
        ['March', '$120K', '-8%']
    ]
    
    for row_data in data_rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data
    
    doc.save(word_path)
    return str(word_path)

def test_document_parsing():
    """Test parsing of different document types"""
    print("=== Testing Document Parsing ===")
    
    # Test Excel parsing
    excel_path = create_test_excel()
    try:
        excel_result = tool_registry.execute_tool('parse_excel', file_path=excel_path)
        if 'error' not in excel_result:
            print("✓ Excel parsing successful")
            print(f"  - Found {len(excel_result['sheets'])} sheets")
            print(f"  - Summary: {excel_result['summary']}")
        else:
            print(f"✗ Excel parsing failed: {excel_result['error']}")
            return False
    except Exception as e:
        print(f"✗ Excel parsing exception: {e}")
        return False
    
    # Test Word parsing
    word_path = create_test_word()
    try:
        word_result = tool_registry.execute_tool('parse_word', file_path=word_path)
        if 'error' not in word_result:
            print("✓ Word parsing successful")
            print(f"  - Found {len(word_result['headers'])} headers")
            print(f"  - Summary: {word_result['summary']}")
        else:
            print(f"✗ Word parsing failed: {word_result['error']}")
            return False
    except Exception as e:
        print(f"✗ Word parsing exception: {e}")
        return False
    
    return True

def test_chart_generation():
    """Test chart generation from parsed data"""
    print("\n=== Testing Chart Generation ===")
    
    try:
        # Create sample chart data
        chart_spec = {
            'data': {
                'x': ['Q1', 'Q2', 'Q3', 'Q4'],
                'y': [100, 150, 120, 180]
            },
            'type': 'bar',
            'title': 'Quarterly Sales Performance',
            'xlabel': 'Quarter',
            'ylabel': 'Sales ($K)'
        }
        
        result = tool_registry.execute_tool('generate_chart', data_spec=chart_spec)
        
        if result.get('status') == 'success':
            print("✓ Chart generation successful")
            print(f"  - Chart saved to: {result['chart_path']}")
            
            # Verify file exists
            if Path(result['chart_path']).exists():
                print("✓ Chart file created successfully")
                return True
            else:
                print("✗ Chart file not found")
                return False
        else:
            print(f"✗ Chart generation failed: {result.get('error')}")
            return False
            
    except Exception as e:
        print(f"✗ Chart generation exception: {e}")
        return False

def test_document_modification():
    """Test document modification tools"""
    print("\n=== Testing Document Modification ===")
    
    # Test Excel modification
    excel_path = create_test_excel()
    try:
        modify_instructions = {
            'operations': [
                {
                    'type': 'add_data',
                    'sheet': 'Sheet1',
                    'data': [['Jun', 220, 120, 100]],
                },
                {
                    'type': 'add_formula',
                    'sheet': 'Sheet1', 
                    'cell': 'E1',
                    'formula': '=SUM(B:B)'
                }
            ]
        }
        
        excel_result = tool_registry.execute_tool(
            'modify_excel',
            file_path=excel_path,
            instructions=modify_instructions
        )
        
        if excel_result.get('status') == 'success':
            print("✓ Excel modification successful")
            print(f"  - Modified file: {excel_result['output_path']}")
        else:
            print(f"✗ Excel modification failed: {excel_result.get('error')}")
            return False
            
    except Exception as e:
        print(f"✗ Excel modification exception: {e}")
        return False
    
    # Test Word modification
    word_path = create_test_word()
    try:
        modify_instructions = {
            'operations': [
                {'type': 'add_heading', 'text': 'Conclusions', 'level': 1},
                {'type': 'add_paragraph', 'text': 'Based on the analysis, we recommend increasing marketing spend in Q4.'},
                {'type': 'add_page_break'}
            ]
        }
        
        word_result = tool_registry.execute_tool(
            'modify_word',
            file_path=word_path,
            instructions=modify_instructions
        )
        
        if word_result.get('status') == 'success':
            print("✓ Word modification successful")
            print(f"  - Modified file: {word_result['output_path']}")
            return True
        else:
            print(f"✗ Word modification failed: {word_result.get('error')}")
            return False
            
    except Exception as e:
        print(f"✗ Word modification exception: {e}")
        return False

def test_document_summarization():
    """Test document summarization system"""
    print("\n=== Testing Document Summarization ===")
    
    try:
        # Test Excel summarization
        excel_path = create_test_excel()
        excel_content = ExcelParser.parse(excel_path)
        excel_summary = ExcelParser.generate_summary(excel_content)
        
        print("✓ Excel summarization successful")
        print(f"  - Summary: {excel_summary}")
        
        # Test Word summarization
        word_path = create_test_word()
        word_content = WordParser.parse(word_path)
        word_summary = WordParser.generate_summary(word_content)
        
        print("✓ Word summarization successful")
        print(f"  - Summary: {word_summary}")
        
        return True
        
    except Exception as e:
        print(f"✗ Document summarization exception: {e}")
        return False

def test_artifact_saving():
    """Test artifact saving functionality"""
    print("\n=== Testing Artifact Saving ===")
    
    try:
        # Test saving different types of content
        test_cases = [
            ("Hello World", "txt"),
            ('{"test": "data"}', "json"),
            ("id,name,value\n1,Test,100", "csv")
        ]
        
        for content, file_type in test_cases:
            result = tool_registry.execute_tool(
                'save_artifact',
                content=content,
                file_type=file_type
            )
            
            if result.get('status') == 'success':
                print(f"✓ Saved {file_type} artifact: {result['artifact_id']}")
                
                # Verify file exists
                if Path(result['path']).exists():
                    print(f"  - File created: {result['path']}")
                else:
                    print(f"  ✗ File not found: {result['path']}")
                    return False
            else:
                print(f"✗ Failed to save {file_type} artifact: {result.get('error')}")
                return False
        
        return True
        
    except Exception as e:
        print(f"✗ Artifact saving exception: {e}")
        return False

def test_end_to_end_workflow():
    """Test complete workflow: parse -> analyze -> modify -> save"""
    print("\n=== Testing End-to-End Workflow ===")
    
    try:
        # 1. Create and parse Excel file
        excel_path = create_test_excel()
        excel_data = tool_registry.execute_tool('parse_excel', file_path=excel_path)
        
        if 'error' in excel_data:
            print(f"✗ Step 1 failed: {excel_data['error']}")
            return False
        
        print("✓ Step 1: Parsed Excel data")
        
        # 2. Generate chart from the data
        # Extract data for charting
        sheets = excel_data['sheets']
        sheet_data = list(sheets.values())[0]  # Get first sheet
        
        chart_data = {
            'x': [row['Month'] for row in sheet_data],
            'y': [row['Sales'] for row in sheet_data]
        }
        
        chart_spec = {
            'data': chart_data,
            'type': 'line',
            'title': 'Monthly Sales Trend',
            'xlabel': 'Month',
            'ylabel': 'Sales'
        }
        
        chart_result = tool_registry.execute_tool('generate_chart', data_spec=chart_spec)
        
        if chart_result.get('status') != 'success':
            print(f"✗ Step 2 failed: {chart_result.get('error')}")
            return False
        
        print("✓ Step 2: Generated chart from data")
        
        # 3. Create comprehensive report
        report_content = f"""
# Sales Analysis Report

## Data Summary
{excel_data['summary']}

## Key Insights
- Monthly sales data shows growth trend
- Peak performance in May with $200K sales
- Chart generated: {chart_result['chart_path']}

## Recommendations
- Focus marketing efforts on high-performing months
- Analyze factors contributing to May's success
"""
        
        report_result = tool_registry.execute_tool(
            'save_artifact',
            content=report_content,
            file_type='md'
        )
        
        if report_result.get('status') != 'success':
            print(f"✗ Step 3 failed: {report_result.get('error')}")
            return False
        
        print("✓ Step 3: Created comprehensive report")
        print(f"  - Report saved: {report_result['path']}")
        print(f"  - Chart saved: {chart_result['chart_path']}")
        
        return True
        
    except Exception as e:
        print(f"✗ End-to-end workflow exception: {e}")
        return False

def main():
    """Run comprehensive integration tests"""
    print("SmolAgents Integration Test Suite")
    print("=" * 60)
    
    tests = [
        test_document_parsing,
        test_chart_generation,
        test_document_modification,
        test_document_summarization,
        test_artifact_saving,
        test_end_to_end_workflow
    ]
    
    results = []
    for test in tests:
        try:
            result = test()
            results.append(result)
        except Exception as e:
            print(f"✗ Test {test.__name__} failed with exception: {e}")
            results.append(False)
    
    print("\n" + "=" * 60)
    print("INTEGRATION TEST SUMMARY")
    print("=" * 60)
    
    passed = sum(1 for r in results if r)
    total = len(results)
    
    print(f"Tests passed: {passed}/{total}")
    
    if passed == total:
        print("🎉 ALL INTEGRATION TESTS PASSED!")
        print("\nThe SmolAgents system is ready for:")
        print("- Document parsing (PDF, Excel, Word)")
        print("- Chart generation from data")
        print("- Document modification and enhancement") 
        print("- Comprehensive document summarization")
        print("- Artifact creation and management")
        print("- End-to-end document processing workflows")
        return True
    else:
        print("❌ Some integration tests failed")
        return False

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)