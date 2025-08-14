"""
Utility functions for SmolAgents tools to improve robustness and error handling.
"""

import json
import re
import logging
from typing import Dict, Any, Union, Optional, List
from pathlib import Path
import jsonschema
from jsonschema import validate, ValidationError

logger = logging.getLogger(__name__)

class ToolInputSanitizer:
    """Handles robust JSON input parsing and validation for tools."""
    
    @staticmethod
    def sanitize_json_input(input_string: str) -> str:
        """
        Sanitize JSON input to handle various quote formats and common issues.
        
        Args:
            input_string (str): Raw JSON input string
            
        Returns:
            str: Sanitized JSON string ready for parsing
        """
        if not isinstance(input_string, str):
            return str(input_string)
        
        # Remove leading/trailing whitespace
        cleaned = input_string.strip()
        
        # Handle empty or null inputs
        if not cleaned or cleaned.lower() in ['null', 'none', '']:
            return '{}'
        
        # If it's already a valid JSON string, return as-is
        try:
            json.loads(cleaned)
            return cleaned
        except json.JSONDecodeError:
            pass
        
        # Handle common quote issues
        # Replace smart quotes with regular quotes
        cleaned = cleaned.replace('"', '"').replace('"', '"')
        cleaned = cleaned.replace(''', "'").replace(''', "'")
        
        # Handle mixed quote scenarios
        # If string starts and ends with single quotes, try to convert internal doubles
        if cleaned.startswith("'") and cleaned.endswith("'") and cleaned.count("'") >= 2:
            # Extract content between outer single quotes
            inner_content = cleaned[1:-1]
            
            # Replace internal single quotes that might conflict
            # This is a simplified approach - in a real scenario, proper JSON parsing would be better
            if '"' in inner_content:
                # If there are double quotes inside, likely the outer singles are the string delimiters
                # Just replace the outer quotes
                cleaned = '"' + inner_content + '"'
        
        # Handle Python-style single quotes in JSON
        # This is a more aggressive approach for when we know it's meant to be JSON
        if "'" in cleaned and cleaned.count("'") > cleaned.count('"'):
            # Try to convert Python dict format to JSON
            try:
                # Replace single quotes with double quotes, but be careful with nested quotes
                cleaned = re.sub(r"'([^']*)'(?=\s*:)", r'"\1"', cleaned)  # Keys
                cleaned = re.sub(r":\s*'([^']*)'", r': "\1"', cleaned)   # String values
                cleaned = re.sub(r"\[\s*'([^']*)'", r'["\1"', cleaned)   # Array string elements
                cleaned = re.sub(r",\s*'([^']*)'", r', "\1"', cleaned)   # Array string elements
            except Exception as e:
                logger.warning(f"Failed to convert Python-style quotes: {e}")
        
        return cleaned
    
    @staticmethod
    def safe_json_loads(input_data: Union[str, Dict, List], fallback_value: Any = None) -> Any:
        """
        Safely parse JSON input with multiple fallback strategies.
        
        Args:
            input_data: Input to parse (string, dict, or list)
            fallback_value: Value to return if all parsing fails
            
        Returns:
            Parsed data or fallback value
        """
        # If already parsed, return as-is
        if isinstance(input_data, (dict, list)):
            return input_data
        
        if not isinstance(input_data, str):
            input_data = str(input_data)
        
        # Strategy 1: Direct JSON loads
        try:
            return json.loads(input_data)
        except json.JSONDecodeError as e:
            logger.debug(f"Direct JSON parsing failed: {e}")
        
        # Strategy 2: Sanitize then parse
        try:
            sanitized = ToolInputSanitizer.sanitize_json_input(input_data)
            return json.loads(sanitized)
        except json.JSONDecodeError as e:
            logger.debug(f"Sanitized JSON parsing failed: {e}")
        
        # Strategy 3: Try eval for Python-style dicts (DANGEROUS - only for trusted input)
        try:
            # Only attempt eval if it looks like a Python dict/list
            if input_data.strip().startswith(('{', '[')):
                result = eval(input_data)
                if isinstance(result, (dict, list)):
                    return result
        except Exception as e:
            logger.debug(f"Python eval parsing failed: {e}")
        
        # Strategy 4: Try to fix common JSON issues
        try:
            # Replace True/False/None with JSON equivalents
            fixed = input_data.replace('True', 'true').replace('False', 'false').replace('None', 'null')
            return json.loads(fixed)
        except json.JSONDecodeError as e:
            logger.debug(f"Fixed JSON parsing failed: {e}")
        
        logger.warning(f"All JSON parsing strategies failed for input: {input_data[:100]}...")
        return fallback_value if fallback_value is not None else {}

class ToolValidator:
    """Validation schemas and methods for tool inputs."""
    
    # Schema definitions for each tool
    SCHEMAS = {
        'excel_generator': {
            'type': 'object',
            'properties': {
                'sheets': {
                    'type': 'array',
                    'items': {
                        'type': 'object',
                        'properties': {
                            'name': {'type': 'string'},
                            'title': {'type': 'string'},
                            'tables': {
                                'type': 'array',
                                'items': {
                                    'type': 'object',
                                    'properties': {
                                        'data': {
                                            'type': 'array',
                                            'items': {'type': 'array'}
                                        },
                                        'headers': {'type': 'array'},
                                        'title': {'type': 'string'}
                                    }
                                }
                            },
                            'charts': {'type': 'array'}
                        },
                        'required': ['name']
                    }
                }
            },
            'required': ['sheets']
        },
        
        'chart_generator': {
            'type': 'object',
            'properties': {
                'type': {
                    'type': 'string',
                    'enum': ['bar', 'line', 'pie', 'scatter']
                },
                'title': {'type': 'string'},
                'data': {
                    'type': 'object',
                    'oneOf': [
                        {
                            'properties': {
                                'x': {'type': 'array'},
                                'y': {'type': 'array'}
                            },
                            'required': ['x', 'y']
                        },
                        {
                            'properties': {
                                'labels': {'type': 'array'},
                                'values': {'type': 'array'}
                            },
                            'required': ['labels', 'values']
                        }
                    ]
                },
                'xlabel': {'type': 'string'},
                'ylabel': {'type': 'string'},
                'save_path': {'type': 'string'}
            },
            'required': ['type', 'data']
        },
        
        'excel_modifier': {
            'type': 'object',
            'anyOf': [
                {
                    'properties': {
                        'operations': {
                            'type': 'array',
                            'items': {
                                'type': 'object',
                                'properties': {
                                    'type': {'type': 'string'},
                                    'name': {'type': 'string'},
                                    'sheet': {'type': 'string'},
                                    'data': {'type': 'array'}
                                },
                                'required': ['type']
                            }
                        }
                    },
                    'required': ['operations']
                },
                {
                    'properties': {
                        'add_sheets': {'type': 'array'},
                        'add_data': {'type': 'array'},
                        'add_charts': {'type': 'array'}
                    }
                }
            ]
        }
    }
    
    @staticmethod
    def validate_input(data: Dict[str, Any], schema_name: str) -> tuple[bool, Optional[str]]:
        """
        Validate input data against a schema.
        
        Args:
            data: Data to validate
            schema_name: Name of the schema to use
            
        Returns:
            tuple: (is_valid, error_message)
        """
        if schema_name not in ToolValidator.SCHEMAS:
            return True, None  # Skip validation if schema not found
        
        try:
            validate(instance=data, schema=ToolValidator.SCHEMAS[schema_name])
            return True, None
        except ValidationError as e:
            return False, f"Validation error: {e.message}"
        except Exception as e:
            return False, f"Validation failed: {str(e)}"
    
    @staticmethod
    def get_example_for_schema(schema_name: str) -> str:
        """Get an example JSON string for a given schema."""
        examples = {
            'excel_generator': '''{
    "sheets": [
        {
            "name": "Sales Data",
            "title": "Monthly Sales Report",
            "tables": [
                {
                    "title": "Sales by Product",
                    "data": [
                        ["Product", "Sales", "Region"],
                        ["Widget A", 1000, "North"],
                        ["Widget B", 1500, "South"]
                    ]
                }
            ]
        }
    ]
}''',
            'chart_generator': '''{
    "type": "bar",
    "title": "Sales Chart",
    "data": {
        "x": ["Product A", "Product B", "Product C"],
        "y": [100, 150, 120]
    },
    "xlabel": "Products",
    "ylabel": "Sales"
}''',
            'excel_modifier': '''{
    "operations": [
        {
            "type": "add_sheet",
            "name": "New Sheet",
            "data": [
                ["Header1", "Header2"],
                ["Value1", "Value2"]
            ]
        }
    ]
}'''
        }
        return examples.get(schema_name, '{}')

class FileVerifier:
    """Utilities for verifying generated files."""
    
    @staticmethod
    def verify_excel_file(file_path: str) -> Dict[str, Any]:
        """
        Verify that an Excel file was created correctly and contains data.
        
        Args:
            file_path: Path to the Excel file
            
        Returns:
            Dict with verification results
        """
        result = {
            'exists': False,
            'size': 0,
            'sheets': [],
            'has_data': False,
            'errors': []
        }
        
        try:
            path = Path(file_path)
            
            # Check if file exists
            if not path.exists():
                result['errors'].append(f"File does not exist: {file_path}")
                return result
            
            result['exists'] = True
            result['size'] = path.stat().st_size
            
            # Check file size
            if result['size'] == 0:
                result['errors'].append("File is empty")
                return result
            
            # Try to read with openpyxl
            try:
                import openpyxl
                wb = openpyxl.load_workbook(file_path, read_only=True)
                result['sheets'] = wb.sheetnames
                
                # Check if sheets have data
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    if ws.max_row > 1 or (ws.max_row == 1 and any(cell.value for cell in ws[1])):
                        result['has_data'] = True
                        break
                
                wb.close()
                
            except Exception as e:
                result['errors'].append(f"Could not read Excel file: {str(e)}")
            
        except Exception as e:
            result['errors'].append(f"Verification failed: {str(e)}")
        
        return result
    
    @staticmethod
    def verify_word_file(file_path: str) -> Dict[str, Any]:
        """
        Verify that a Word file was created correctly and contains content.
        
        Args:
            file_path: Path to the Word file
            
        Returns:
            Dict with verification results
        """
        result = {
            'exists': False,
            'size': 0,
            'has_content': False,
            'errors': []
        }
        
        try:
            path = Path(file_path)
            
            # Check if file exists
            if not path.exists():
                result['errors'].append(f"File does not exist: {file_path}")
                return result
            
            result['exists'] = True
            result['size'] = path.stat().st_size
            
            # Check file size
            if result['size'] == 0:
                result['errors'].append("File is empty")
                return result
            
            # Try to read with python-docx
            try:
                from docx import Document
                doc = Document(file_path)
                
                # Check if document has content
                if len(doc.paragraphs) > 0:
                    for para in doc.paragraphs:
                        if para.text.strip():
                            result['has_content'] = True
                            break
                
            except Exception as e:
                result['errors'].append(f"Could not read Word file: {str(e)}")
            
        except Exception as e:
            result['errors'].append(f"Verification failed: {str(e)}")
        
        return result

class DebugLogger:
    """Enhanced logging for debugging tool operations."""
    
    @staticmethod
    def log_tool_start(tool_name: str, inputs: Dict[str, Any]):
        """Log the start of tool operation."""
        logger.info(f"[{tool_name}] Starting tool execution")
        logger.debug(f"[{tool_name}] Inputs: {str(inputs)[:200]}...")
    
    @staticmethod
    def log_json_parsing(tool_name: str, raw_input: str, parsed_result: Any):
        """Log JSON parsing details."""
        logger.debug(f"[{tool_name}] JSON parsing - Input length: {len(raw_input)}")
        logger.debug(f"[{tool_name}] JSON parsing - Result type: {type(parsed_result)}")
        if isinstance(parsed_result, dict):
            logger.debug(f"[{tool_name}] JSON parsing - Keys: {list(parsed_result.keys())}")
    
    @staticmethod
    def log_file_operation(tool_name: str, operation: str, file_path: str, success: bool, details: str = ""):
        """Log file operations."""
        status = "SUCCESS" if success else "FAILED"
        logger.info(f"[{tool_name}] File {operation} {status}: {file_path}")
        if details:
            logger.debug(f"[{tool_name}] Details: {details}")
    
    @staticmethod
    def log_validation_result(tool_name: str, is_valid: bool, error_msg: str = ""):
        """Log validation results."""
        status = "PASSED" if is_valid else "FAILED"
        logger.info(f"[{tool_name}] Input validation {status}")
        if error_msg:
            logger.warning(f"[{tool_name}] Validation error: {error_msg}")

class ErrorFormatter:
    """Format user-friendly error messages with examples."""
    
    @staticmethod
    def format_json_error(tool_name: str, raw_input: str, error: Exception) -> str:
        """Format JSON parsing error with helpful guidance."""
        example = ToolValidator.get_example_for_schema(tool_name.replace('_tool', ''))
        
        return f"""
JSON parsing failed for {tool_name}:
Error: {str(error)}

Input received: {raw_input[:100]}...

Expected format example:
{example}

Common fixes:
1. Use double quotes for JSON keys and string values
2. Ensure proper escaping of quotes inside strings
3. Make sure brackets and braces are properly matched
4. Check for trailing commas (not allowed in JSON)
"""
    
    @staticmethod
    def format_validation_error(tool_name: str, validation_error: str) -> str:
        """Format validation error with schema guidance."""
        example = ToolValidator.get_example_for_schema(tool_name.replace('_tool', ''))
        
        return f"""
Input validation failed for {tool_name}:
{validation_error}

Expected format example:
{example}

Please check your input structure matches the required schema.
"""
    
    @staticmethod
    def format_file_creation_error(tool_name: str, file_path: str, verification_result: Dict[str, Any]) -> str:
        """Format file creation error with verification details."""
        errors = verification_result.get('errors', [])
        
        message = f"File creation issue for {tool_name}:\n"
        message += f"Target file: {file_path}\n"
        
        if verification_result.get('exists'):
            message += f"File exists but may have issues:\n"
            message += f"- Size: {verification_result.get('size', 0)} bytes\n"
            if 'sheets' in verification_result:
                message += f"- Sheets: {verification_result['sheets']}\n"
            message += f"- Has data: {verification_result.get('has_data', False)}\n"
        else:
            message += "File was not created.\n"
        
        if errors:
            message += f"Errors found:\n"
            for error in errors:
                message += f"- {error}\n"
        
        return message

class ToolOutputSanitizer:
    """Sanitize tool outputs to prevent framework parsing issues."""
    
    @staticmethod
    def sanitize_tool_output(output: Any, tool_name: str) -> Any:
        """
        Sanitize tool output to prevent SmolAgents from misinterpreting it as a tool definition.
        
        Args:
            output: Raw tool output
            tool_name: Name of the tool that generated the output
            
        Returns:
            Sanitized output safe for SmolAgents framework
        """
        try:
            # If output is a dictionary that might be confused with tool definition
            if isinstance(output, dict):
                # Check if it looks like a tool definition (has keys that could confuse the parser)
                suspicious_keys = {'name', 'description', 'inputs', 'output_type', 'type', 'title', 'data'}
                output_keys = set(output.keys())
                
                # If output has suspicious keys that might confuse the parser
                if suspicious_keys.intersection(output_keys):
                    # Wrap it in a clear structure that identifies it as tool output
                    sanitized = {
                        '_tool_output': True,
                        '_source_tool': tool_name,
                        '_output_data': output,
                        '_timestamp': str(logger.handlers[0].formatter.formatTime(logger.makeRecord('', 0, '', 0, '', (), None)) if logger.handlers else 'unknown')
                    }
                    logger.debug(f"Sanitized potentially confusing output from {tool_name}")
                    return sanitized
            
            # For string outputs, ensure they don't look like JSON tool definitions
            elif isinstance(output, str):
                # Check if string looks like JSON that could be misinterpreted
                if output.strip().startswith('{') and 'name' in output and 'description' in output:
                    # Wrap in safe structure
                    sanitized = {
                        '_tool_output': True,
                        '_source_tool': tool_name,
                        '_output_text': output,
                        '_timestamp': 'generated'
                    }
                    logger.debug(f"Sanitized potentially confusing string output from {tool_name}")
                    return sanitized
            
            # Output looks safe, return as-is
            return output
            
        except Exception as e:
            logger.warning(f"Error sanitizing output from {tool_name}: {e}")
            # If sanitization fails, return a safe wrapper
            return {
                '_tool_output': True,
                '_source_tool': tool_name,
                '_raw_output': str(output),
                '_sanitization_error': str(e),
                '_timestamp': 'error'
            }
    
    @staticmethod
    def extract_actual_output(sanitized_output: Any) -> Any:
        """
        Extract the actual output from a sanitized wrapper.
        
        Args:
            sanitized_output: Potentially sanitized output
            
        Returns:
            Original output data
        """
        if isinstance(sanitized_output, dict) and sanitized_output.get('_tool_output'):
            # Extract from sanitized wrapper
            if '_output_data' in sanitized_output:
                return sanitized_output['_output_data']
            elif '_output_text' in sanitized_output:
                return sanitized_output['_output_text']
            elif '_raw_output' in sanitized_output:
                return sanitized_output['_raw_output']
        
        # Not sanitized, return as-is
        return sanitized_output