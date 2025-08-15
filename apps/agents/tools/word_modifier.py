from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
from typing import Dict, Any, List, Optional
from pathlib import Path
from PIL import Image
import io

logger = logging.getLogger(__name__)

# Import preview generator
try:
    from .word_preview import WordPreviewGenerator
except ImportError:
    logger.warning("WordPreviewGenerator not available - previews will not be generated")
    WordPreviewGenerator = None

class WordModifier:
    """Modify Word documents and insert images/charts"""
    
    @staticmethod
    def modify_word(
        file_path: str,
        instructions: Dict[str, Any],
        images: List[str] = None,
        output_path: Optional[str] = None
    ) -> str:
        """
        Modify Word document based on instructions
        
        Supports multiple instruction formats:
        1. Operations format:
        {
            'operations': [
                {'type': 'add_heading', 'text': 'New Section', 'level': 1},
                {'type': 'add_paragraph', 'text': 'Some text...'},
                {'type': 'add_table', 'data': [[...]], 'headers': [...]},
            ]
        }
        
        2. Direct format:
        {
            'add_headings': [{'text': 'Title', 'level': 1}],
            'add_paragraphs': [{'text': 'Content'}],
            'add_tables': [{'data': [[...]], 'style': 'Table Grid'}]
        }
        """
        try:
            # Open existing document or create new
            if file_path and Path(file_path).exists():
                doc = Document(file_path)
            else:
                doc = Document()
            
            # Convert instructions to operations format if needed
            operations = WordModifier._normalize_instructions(instructions)
            
            # Process operations
            for operation in operations:
                op_type = operation['type']
                
                if op_type == 'add_heading':
                    doc.add_heading(operation['text'], level=operation.get('level', 1))
                
                elif op_type == 'add_paragraph':
                    paragraph = doc.add_paragraph(operation['text'])
                    if operation.get('bold'):
                        paragraph.runs[0].bold = True
                    if operation.get('italic'):
                        paragraph.runs[0].italic = True
                    if operation.get('alignment'):
                        alignment_map = {
                            'center': WD_ALIGN_PARAGRAPH.CENTER,
                            'right': WD_ALIGN_PARAGRAPH.RIGHT,
                            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
                        }
                        paragraph.alignment = alignment_map.get(operation['alignment'])
                
                elif op_type == 'add_table':
                    WordModifier._add_table(doc, operation)
                
                elif op_type == 'insert_image' and images:
                    image_index = operation.get('image_index', 0)
                    if image_index < len(images):
                        width = Inches(operation.get('width', 6))
                        doc.add_picture(images[image_index], width=width)
                
                elif op_type == 'add_page_break':
                    doc.add_page_break()
            
            # Save document
            if not output_path:
                # Use cross-platform temp directory within project
                outputs_dir = Path.cwd() / 'temp' / 'outputs'
                outputs_dir.mkdir(parents=True, exist_ok=True)
                
                if file_path and Path(file_path).exists():
                    output_path = outputs_dir / f"modified_{Path(file_path).stem}.docx"
                else:
                    output_path = outputs_dir / "new_document.docx"
            
            doc.save(output_path)
            
            # Generate HTML preview if possible
            preview_html = None
            if WordPreviewGenerator:
                try:
                    preview_result = WordPreviewGenerator.generate_preview(str(output_path))
                    if preview_result['success']:
                        preview_html = preview_result['preview_html']
                        logger.info(f"Generated HTML preview for modified document: {output_path}")
                    else:
                        logger.warning(f"Failed to generate preview for {output_path}: {preview_result.get('error', 'Unknown error')}")
                except Exception as e:
                    logger.error(f"Error generating preview for {output_path}: {str(e)}")
            
            # Return structured result including preview
            result = {
                'file_path': str(output_path),
                'preview_html': preview_html,
                'message': f"Word document modified successfully: {output_path}"
            }
            
            return str(result)  # Convert to string for tool compatibility
            
        except Exception as e:
            logger.error(f"Error modifying Word document: {str(e)}")
            raise
    
    @staticmethod
    def _normalize_instructions(instructions: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Convert various instruction formats to normalized operations list
        """
        operations = []
        
        # If already in operations format, return as-is
        if 'operations' in instructions:
            return instructions['operations']
        
        # Handle direct format
        # Add headings
        if 'add_headings' in instructions:
            for heading in instructions['add_headings']:
                operations.append({
                    'type': 'add_heading',
                    'text': heading['text'],
                    'level': heading.get('level', 1)
                })
        
        # Add paragraphs
        if 'add_paragraphs' in instructions:
            for paragraph in instructions['add_paragraphs']:
                operations.append({
                    'type': 'add_paragraph',
                    'text': paragraph['text'],
                    'bold': paragraph.get('bold', False),
                    'italic': paragraph.get('italic', False),
                    'alignment': paragraph.get('alignment')
                })
        
        # Add tables
        if 'add_tables' in instructions:
            for table in instructions['add_tables']:
                operations.append({
                    'type': 'add_table',
                    'data': table['data'],
                    'headers': table.get('headers'),
                    'style': table.get('style', 'Table Grid')
                })
        
        # Handle legacy single operations
        if 'add_paragraph' in instructions:
            operations.append({
                'type': 'add_paragraph',
                'text': instructions['add_paragraph']
            })
        
        if 'add_heading' in instructions:
            operations.append({
                'type': 'add_heading',
                'text': instructions['add_heading'],
                'level': instructions.get('heading_level', 1)
            })
        
        if 'add_table' in instructions:
            table_config = instructions['add_table']
            operations.append({
                'type': 'add_table',
                'data': table_config.get('data', []),
                'headers': table_config.get('headers'),
                'style': table_config.get('style', 'Table Grid')
            })
        
        return operations
    
    @staticmethod
    def _add_table(doc: Document, config: Dict[str, Any]):
        """Add table to document"""
        data = config['data']
        headers = config.get('headers', [])
        style = config.get('style', 'Table Grid')
        
        # Create table - if no headers but data has headers as first row
        if not headers and data:
            # Check if first row looks like headers (assume it is if no explicit headers)
            headers = data[0]
            data = data[1:] if len(data) > 1 else []
        
        # Create table
        if data:
            table = doc.add_table(rows=1 if headers else 0, cols=len(data[0]) if data else len(headers))
        else:
            table = doc.add_table(rows=1, cols=len(headers))
        
        table.style = style
        
        # Add headers
        if headers:
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                if i < len(header_cells):
                    header_cells[i].text = str(header)
                    if header_cells[i].paragraphs:
                        for run in header_cells[i].paragraphs[0].runs:
                            run.bold = True
        
        # Add data
        for row_data in data:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                if i < len(row_cells):
                    row_cells[i].text = str(cell_data)