"""
Word Generator Tool for Smol Agents

This tool allows the agent to create Word documents (.docx) with formatted text,
tables, and basic styling based on user requests.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import os
import json
import re
import logging
from typing import Dict, List, Any, Optional
from smolagents import Tool

# Import our robust utility functions
from .tool_utils import (
    ToolInputSanitizer, ToolValidator, FileVerifier,
    DebugLogger, ErrorFormatter
)
from .word_preview import WordPreviewGenerator

logger = logging.getLogger(__name__)

class SimpleWordGeneratorTool(Tool):
    """
    Simplified Word generator for quick text document creation.
    """
    
    name = "simple_word_generator"
    description = """
    Generate Word documents from text content with markdown formatting support.
    
    This tool creates professional Word documents from text with markdown-style formatting:
    - Headers (# ## ###)
    - Bold (**text**) and italic (*text*)
    - Lists (- item or 1. item)
    - Code blocks (```code```)
    - Inline code (`code`)
    
    Example usage:
    simple_word_generator(
        title="Report Title",
        content="# Introduction\\n\\nThis is **bold** text with *italic* formatting.",
        filename="my_report"
    )
    """
    
    inputs = {
        "title": {
            "type": "string",
            "description": "Document title"
        },
        "content": {
            "type": "string",
            "description": "Main document content (paragraphs separated by double newlines), writed in markdown style"
        },
        "filename": {
            "type": "string",
            "description": "Name for the Word document (without extension)"
        }
    }
    
    output_type = "string"
    
    def _parse_markdown_line(self, line: str, doc: Document):
        """Parse a single line of markdown and add appropriate formatting to the document."""
        line = line.strip()
        if not line:
            return
        
        # Headers
        if line.startswith('#'):
            level = 0
            while level < len(line) and line[level] == '#':
                level += 1
            if level <= 6 and level < len(line) and line[level] == ' ':
                header_text = line[level + 1:].strip()
                doc.add_heading(header_text, level)
                return
        
        # Lists
        if line.startswith('- ') or line.startswith('* '):
            list_text = line[2:].strip()
            paragraph = doc.add_paragraph()
            paragraph.style = 'List Bullet'
            self._add_formatted_text(paragraph, list_text)
            return
        
        # Numbered lists
        if re.match(r'^\d+\.\s', line):
            list_text = re.sub(r'^\d+\.\s', '', line)
            paragraph = doc.add_paragraph()
            paragraph.style = 'List Number'
            self._add_formatted_text(paragraph, list_text)
            return
        
        # Code blocks
        if line.startswith('```'):
            return 'code_block'
        
        # Regular paragraph
        paragraph = doc.add_paragraph()
        self._add_formatted_text(paragraph, line)
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add text with inline formatting (bold, italic, code) to a paragraph."""
        # Handle inline code first
        parts = re.split(r'(`[^`]+`)', text)
        
        for part in parts:
            if part.startswith('`') and part.endswith('`'):
                # Inline code
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            else:
                # Handle bold and italic
                self._add_bold_italic_text(paragraph, part)
    
    def _add_bold_italic_text(self, paragraph, text: str):
        """Add text with bold and italic formatting to a paragraph."""
        # Split by bold (**text**)
        bold_parts = re.split(r'(\*\*[^*]+\*\*)', text)
        
        for bold_part in bold_parts:
            if bold_part.startswith('**') and bold_part.endswith('**'):
                # Bold text
                bold_text = bold_part[2:-2]
                # Check for italic within bold
                italic_parts = re.split(r'(\*[^*]+\*)', bold_text)
                for italic_part in italic_parts:
                    if italic_part.startswith('*') and italic_part.endswith('*'):
                        run = paragraph.add_run(italic_part[1:-1])
                        run.bold = True
                        run.italic = True
                    else:
                        run = paragraph.add_run(italic_part)
                        run.bold = True
            else:
                # Check for italic only
                italic_parts = re.split(r'(\*[^*]+\*)', bold_part)
                for italic_part in italic_parts:
                    if italic_part.startswith('*') and italic_part.endswith('*'):
                        run = paragraph.add_run(italic_part[1:-1])
                        run.italic = True
                    else:
                        run = paragraph.add_run(italic_part)

    def forward(self, title: str, content: str, filename: str) -> str:
        """Generate a Word document from markdown content with robust error handling."""
        
        # Initialize debug logging
        DebugLogger.log_tool_start('simple_word_generator', {
            'title': title,
            'content_length': len(str(content)),
            'filename': filename
        })
        
        try:
            # Validate inputs
            if not title and not content:
                return "Error: Both title and content cannot be empty"
            
            if not filename:
                filename = "document"
            
            # Sanitize inputs
            title = str(title) if title else ""
            content = str(content) if content else ""
            filename = str(filename) if filename else "document"
            
            # Create document
            doc = Document()
            
            # Add title
            if title.strip():
                title_heading = doc.add_heading(title.strip(), 0)
                title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logger.debug(f"Added title: {title}")
            
            # Parse markdown content
            if content.strip():
                lines = content.split('\n')
                in_code_block = False
                code_block_content = []
                
                i = 0
                while i < len(lines):
                    line = lines[i]
                    
                    # Handle code blocks
                    if line.strip().startswith('```'):
                        if not in_code_block:
                            in_code_block = True
                            code_block_content = []
                            # Extract language if specified
                            lang = line.strip()[3:].strip()
                        else:
                            # End of code block
                            in_code_block = False
                            if code_block_content:
                                # Add code block as a paragraph with monospace font
                                code_text = '\n'.join(code_block_content)
                                paragraph = doc.add_paragraph()
                                run = paragraph.add_run(code_text)
                                run.font.name = 'Courier New'
                                run.font.size = Pt(10)
                                paragraph.style = 'No Spacing'
                        i += 1
                        continue
                    
                    if in_code_block:
                        code_block_content.append(line)
                        i += 1
                        continue
                    
                    # Handle empty lines (paragraph breaks)
                    if not line.strip():
                        # Look ahead to see if we need to add spacing
                        if i + 1 < len(lines) and lines[i + 1].strip():
                            doc.add_paragraph()
                        i += 1
                        continue
                    
                    # Parse regular markdown line
                    try:
                        self._parse_markdown_line(line, doc)
                    except Exception as e:
                        logger.warning(f"Error parsing line '{line}': {e}")
                        # Fallback: add as plain text
                        doc.add_paragraph(line)
                    
                    i += 1
            else:
                # Add empty paragraph if no content
                doc.add_paragraph("Document created with no content.")
            
            # Ensure filename has .docx extension
            if not filename.endswith('.docx'):
                filename += '.docx'
            
            # Create temp directory if it doesn't exist
            temp_dir = os.path.join(os.getcwd(), 'temp')
            os.makedirs(temp_dir, exist_ok=True)
            
            # Save file
            output_path = os.path.join(temp_dir, filename)
            doc.save(output_path)
            
            # Verify file creation
            verification = FileVerifier.verify_word_file(output_path)
            DebugLogger.log_file_operation(
                'simple_word_generator',
                'creation',
                output_path,
                verification['exists'] and verification['has_content'],
                f"Size: {verification.get('size', 0)} bytes"
            )
            
            if not verification['exists']:
                return f"Error: Word file was not created: {output_path}"
            
            if not verification['has_content']:
                logger.warning("Word file created but appears to have no content")
                return f"Word document created but may be empty: {output_path}"
            
            # Generate HTML preview
            preview_result = WordPreviewGenerator.generate_preview(output_path)
            preview_html = None
            if preview_result['success']:
                preview_html = preview_result['preview_html']
                logger.info(f"Generated HTML preview for {filename}")
            else:
                logger.warning(f"Failed to generate preview for {filename}: {preview_result.get('error', 'Unknown error')}")
            
            # Return structured result including preview
            result = {
                'file_path': output_path,
                'preview_html': preview_html,
                'message': f"Word document with markdown formatting created successfully: {output_path}"
            }
            
            return str(result)  # Convert to string for tool compatibility
            
        except Exception as e:
            logger.error(f"Error in simple_word_generator: {str(e)}", exc_info=True)
            return f"Error creating Word document with markdown formatting: {str(e)}"