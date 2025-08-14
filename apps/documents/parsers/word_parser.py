import logging
from typing import Dict, List, Any
from dataclasses import dataclass
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)

@dataclass
class WordContent:
    paragraphs: List[str]
    tables: List[List[List[str]]]
    headers: List[str]
    metadata: Dict[str, Any]

class WordParser:
    """Parse Word documents"""
    
    @staticmethod
    def parse(file_path: str) -> WordContent:
        """Parse Word document and extract content"""
        try:
            doc = Document(file_path)
            
            paragraphs = []
            tables = []
            headers = []
            
            # Extract paragraphs and headers
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
                    
                    # Check if it's a heading
                    if para.style.name.startswith('Heading'):
                        headers.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables.append(table_data)
            
            # Extract metadata
            metadata = {
                'paragraph_count': len(paragraphs),
                'table_count': len(tables),
                'header_count': len(headers),
                'word_count': sum(len(p.split()) for p in paragraphs),
            }
            
            # Document properties
            core_props = doc.core_properties
            if core_props:
                metadata.update({
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'created': str(core_props.created) if core_props.created else '',
                    'modified': str(core_props.modified) if core_props.modified else '',
                })
            
            return WordContent(
                paragraphs=paragraphs,
                tables=tables,
                headers=headers,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Error parsing Word document {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def generate_summary(content: WordContent, max_length: int = 500) -> str:
        """Generate a summary of the Word content"""
        summary_parts = []
        
        # Basic info
        summary_parts.append(
            f"Word document with {content.metadata['paragraph_count']} paragraphs, "
            f"{content.metadata['word_count']} words"
        )
        
        # Add title if available
        if content.metadata.get('title'):
            summary_parts.append(f"Title: {content.metadata['title']}")
        
        # Add headers preview
        if content.headers:
            headers_preview = ', '.join(content.headers[:3])
            summary_parts.append(f"Headers: {headers_preview}")
        
        # Add text preview
        if content.paragraphs:
            text_preview = ' '.join(content.paragraphs[:2])[:200]
            summary_parts.append(f"Preview: {text_preview}...")
        
        return ' | '.join(summary_parts)[:max_length]